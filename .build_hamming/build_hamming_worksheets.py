from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DIR = Path("/Users/lupin/work/mudassir/codes_correction/hamming_circle")
STUDENT = OUT_DIR / "Hamming_74_Signal_Rescue_Worksheet.docx"
KEY = OUT_DIR / "Hamming_74_Signal_Rescue_Answer_Key.docx"

BLUE = "3D8DFF"
CYAN = "6DCBF4"
PALE = "D0EDFA"
PALE2 = "E8EEF5"
INK = "0B2545"
MUTED = "59616D"
RULE = "B8BCC4"
CORAL = "FF625E"
CORAL_PALE = "FFE3E1"
GREEN = "158F63"
GREEN_PALE = "DDF5E9"
AMBER = "D88900"
AMBER_PALE = "FFF0C9"
WHITE = "FFFFFF"


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=RULE, size=8):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = tc_borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False


def set_cell_text(cell, text, bold=False, color=INK, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def bottom_border(paragraph, color=RULE, size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep(paragraph, next_=False):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepLines")
    p_pr.append(keep)
    if next_:
        keep_next = OxmlElement("w:keepNext")
        p_pr.append(keep_next)


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc, running_label):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ):
        st = doc.styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.keep_with_next = True

    header = sec.header
    p = header.paragraphs[0]
    p.text = running_label
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.runs[0]
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    bottom_border(p, color=RULE, size=6)
    page_field(sec.footer.paragraphs[0])


def add_title(doc, title, subtitle, key=False, new_page=False, top_spacer=False):
    if new_page and doc.paragraphs:
        previous = doc.sections[-1]
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section_type = previous._sectPr.find(qn("w:type"))
        if section_type is None:
            section_type = OxmlElement("w:type")
            previous._sectPr.insert(0, section_type)
        section_type.set(qn("w:val"), "nextPage")
        for attr in (
            "page_width", "page_height", "top_margin", "bottom_margin",
            "left_margin", "right_margin", "header_distance", "footer_distance",
        ):
            setattr(section, attr, getattr(previous, attr))
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
    if top_spacer:
        spacer = doc.add_paragraph(" ")
        spacer.paragraph_format.space_after = Pt(52)
        spacer.paragraph_format.line_spacing = Pt(6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(25)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(INK)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(subtitle)
    r2.font.name = "Calibri"
    r2.font.size = Pt(12)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor.from_string(GREEN if key else BLUE)
    bottom_border(p2, color=GREEN if key else BLUE, size=14)


def add_identity_line(doc):
    t = doc.add_table(rows=1, cols=3)
    set_table_geometry(t, [3600, 2880, 2880])
    labels = ["Name:", "Team:", "Date:"]
    for i, label in enumerate(labels):
        set_cell_text(t.cell(0, i), label + " __________________", align=WD_ALIGN_PARAGRAPH.LEFT, size=10)
        shade_cell(t.cell(0, i), WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, text, fill=PALE, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, fill)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(color)
    set_keep(p)


def task_heading(doc, label, title, color=BLUE):
    p = doc.add_paragraph(style="Heading 2")
    r1 = p.add_run(label + "  ")
    r1.font.color.rgb = RGBColor.from_string(color)
    r1.bold = True
    r2 = p.add_run(title)
    r2.font.color.rgb = RGBColor.from_string(INK)
    r2.bold = True
    set_keep(p, next_=True)
    return p


def response_lines(doc, n=2, prompt=None):
    if prompt:
        p = doc.add_paragraph(prompt)
        p.paragraph_format.space_after = Pt(3)
    for _ in range(n):
        p = doc.add_paragraph(" ")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        bottom_border(p)


def add_bit_grid(doc, bits=None, roles=True, fixed7=False, color_bad=None):
    bits = bits or ["", "", "", "", "", "", ""]
    rows = 3 if roles else 2
    table = doc.add_table(rows=rows, cols=7)
    set_table_geometry(table, [1337, 1337, 1337, 1337, 1337, 1337, 1338])
    for i in range(7):
        pos = i + 1
        set_cell_text(table.cell(0, i), pos, bold=True, color=MUTED, size=9)
        shade_cell(table.cell(0, i), PALE2)
        check = pos in (1, 2, 4)
        set_cell_text(table.cell(1, i), bits[i], bold=True, color=CORAL if color_bad == pos else INK, size=15)
        shade_cell(table.cell(1, i), CORAL_PALE if color_bad == pos else (PALE if check else WHITE))
        if roles:
            role = {1: "check 1", 2: "check 2", 3: "data 1", 4: "check 4", 5: "data 2", 6: "data 3", 7: "fixed 0" if fixed7 else "data 4"}[pos]
            set_cell_text(table.cell(2, i), role, bold=check, color=BLUE if check else MUTED, size=8)
            shade_cell(table.cell(2, i), WHITE)
    return table


def add_coverage_grid(doc, answer=False):
    table = doc.add_table(rows=4, cols=8)
    set_table_geometry(table, [1200] + [1165] * 7)
    headers = ["position", "1", "2", "3", "4", "5", "6", "7"]
    for c, text in enumerate(headers):
        set_cell_text(table.cell(0, c), text, bold=True, color=INK, size=9)
        shade_cell(table.cell(0, c), PALE2)
    sets = [("check 1", {1, 3, 5, 7}, CYAN), ("check 2", {2, 3, 6, 7}, BLUE), ("check 4", {4, 5, 6, 7}, "7B61FF")]
    for r, (label, active, color) in enumerate(sets, start=1):
        set_cell_text(table.cell(r, 0), label, bold=True, color=color, size=9)
        shade_cell(table.cell(r, 0), WHITE)
        for pos in range(1, 8):
            text = "YES" if answer and pos in active else ""
            set_cell_text(table.cell(r, pos), text, bold=True, color=WHITE if (answer and pos in active) else MUTED, size=9)
            shade_cell(table.cell(r, pos), color if (answer and pos in active) else WHITE)
    set_repeat_table_header(table.rows[0])
    return table


def add_check_math(doc, values=None):
    values = values or ["check 1 count: __________  even? □", "check 2 count: __________  even? □", "check 4 count: __________  even? □"]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    colors = [CYAN, BLUE, "7B61FF"]
    for i, text in enumerate(values):
        set_cell_text(table.cell(0, i), text, bold=True, color=colors[i], size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        shade_cell(table.cell(0, i), WHITE)
    return table


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_student():
    doc = Document()
    configure_document(doc, "SIGNAL RESCUE LAB  |  STUDENT WORKSHEET")
    add_title(doc, "Signal Rescue Lab", "Hamming(7,4) student worksheet | Grades 9-10 | Even parity")
    add_identity_line(doc)
    add_callout(doc, "Finish line: independently encode four message bits into seven, repair any single flipped bit, and recover the original message.")

    task_heading(doc, "1A", "The impossible repair")
    p = doc.add_paragraph("A transmission arrives as:  1 0 1 ? 1 1 0")
    p.runs[0].bold = True
    response_lines(doc, 2, "What was the missing bit? Explain why the receiver can or cannot know.")

    task_heading(doc, "1B", "Even-parity sprint")
    p = doc.add_paragraph("Add one check bit so the total number of 1s is even.")
    t = doc.add_table(rows=5, cols=4)
    set_table_geometry(t, [2600, 1900, 2260, 2600])
    for c, text in enumerate(["data", "check bit", "total 1s", "audit: even?"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=9)
        shade_cell(t.cell(0, c), PALE2)
    for r, data in enumerate(["101", "111", "000", "0101"], start=1):
        set_cell_text(t.cell(r, 0), data, bold=True, size=11)
        for c in range(1, 4):
            set_cell_text(t.cell(r, c), "")
    set_repeat_table_header(t.rows[0])

    task_heading(doc, "1C", "Address detective")
    doc.add_paragraph("Fill YES when a position belongs to a check. Then write the failed-label sum for positions 3, 5, 6, and 7.")
    add_coverage_grid(doc, answer=False)
    response_lines(doc, 1, "Addresses: 3 = __________   5 = __________   6 = __________   7 = __________")

    add_title(doc, "Build the seven-slot machine", "Page 2 | Position roles, check groups, and a three-message-bit warm-up", new_page=True)
    task_heading(doc, "2A", "Label the slots")
    doc.add_paragraph("Write check 1, check 2, data 1, check 4, data 2, data 3, data 4 under positions 1-7.")
    add_bit_grid(doc, bits=["", "", "", "", "", "", ""], roles=False)
    add_callout(doc, "Reference: check 1 covers 1,3,5,7 | check 2 covers 2,3,6,7 | check 4 covers 4,5,6,7", fill=PALE2)

    task_heading(doc, "2B", "Training wheels: encode 1 0 1")
    doc.add_paragraph("Place 1,0,1 in positions 3,5,6. Fix position 7 at 0. Choose each check bit to make its group even.")
    add_bit_grid(doc, bits=["", "", "1", "", "0", "1", "0"], fixed7=True)
    add_check_math(doc)
    response_lines(doc, 1, "Completed seven-bit word: __________________________________________")

    task_heading(doc, "2C", "One more restricted warm-up")
    doc.add_paragraph("Encode message bits 0 1 1 with position 7 fixed at 0.")
    add_bit_grid(doc, bits=["", "", "0", "", "1", "1", "0"], fixed7=True)
    response_lines(doc, 1, "Completed seven-bit word: __________________________________________")
    p = doc.add_paragraph("Accuracy note: this is a restricted practice version of the full Hamming(7,4) code, not the standard Hamming(7,3) code.")
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    add_title(doc, "Encode Hamming(7,4)", "Page 3 | Four message bits become seven transmitted bits", new_page=True)
    task_heading(doc, "3A", "Write the four-move routine")
    t = doc.add_table(rows=2, cols=4)
    set_table_geometry(t, [2340, 2340, 2340, 2340])
    for i, text in enumerate(["1. Label", "2. Place data", "3. Set checks", "4. Verify"]):
        set_cell_text(t.cell(0, i), text, bold=True, color=BLUE if i < 3 else CORAL, size=10)
        shade_cell(t.cell(0, i), PALE2)
        set_cell_text(t.cell(1, i), "", size=9)

    task_heading(doc, "3B", "Encode the message 1 0 1 1")
    add_bit_grid(doc, bits=["", "", "1", "", "0", "1", "1"])
    add_check_math(doc)
    response_lines(doc, 1, "Codeword: ____________________________   Audit complete? □")

    task_heading(doc, "3C", "Speed round")
    doc.add_paragraph("Encode each message. A teammate should audit all three checks.")
    t = doc.add_table(rows=5, cols=4)
    set_table_geometry(t, [1800, 3300, 2300, 1960])
    for c, text in enumerate(["message", "seven-bit codeword", "audit initials", "all even?"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=9)
        shade_cell(t.cell(0, c), PALE2)
    for r, msg in enumerate(["0000", "1111", "0101", "1100"], start=1):
        set_cell_text(t.cell(r, 0), msg, bold=True, size=11)
        set_cell_text(t.cell(r, 1), "")
        set_cell_text(t.cell(r, 2), "")
        set_cell_text(t.cell(r, 3), "□")
    set_repeat_table_header(t.rows[0])
    add_callout(doc, "Encoder shortcut: data always goes in positions 3,5,6,7. Check bits always go in 1,2,4.", fill=GREEN_PALE, color=GREEN)

    add_title(doc, "Find and repair one bad bit", "Page 4 | Run checks, add failed labels, flip once, extract data", new_page=True)
    task_heading(doc, "4A", "Repair 0 1 1 0 0 0 1")
    add_bit_grid(doc, bits=["0", "1", "1", "0", "0", "0", "1"], color_bad=None)
    t = doc.add_table(rows=4, cols=5)
    set_table_geometry(t, [1500, 3000, 1500, 1500, 1860])
    for c, text in enumerate(["check", "positions", "number of 1s", "pass/fail", "failed label"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=8)
        shade_cell(t.cell(0, c), PALE2)
    checks = [("1", "1,3,5,7"), ("2", "2,3,6,7"), ("4", "4,5,6,7")]
    for r, (label, positions) in enumerate(checks, start=1):
        set_cell_text(t.cell(r, 0), label, bold=True)
        set_cell_text(t.cell(r, 1), positions)
        for c in range(2, 5):
            set_cell_text(t.cell(r, c), "")
    response_lines(doc, 2, "Syndrome address: ______   Flip position: ______   Corrected word: __________________\nRecovered message from positions 3,5,6,7: __________________")

    task_heading(doc, "4B", "Packet repair race")
    t = doc.add_table(rows=3, cols=4)
    set_table_geometry(t, [1700, 2500, 2580, 2580])
    for c, text in enumerate(["packet", "received", "error position", "corrected / message"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=9)
        shade_cell(t.cell(0, c), PALE2)
    for r, (name, received) in enumerate([("A", "0101101"), ("B", "0111101")], start=1):
        set_cell_text(t.cell(r, 0), name, bold=True, color=BLUE if r == 1 else CORAL)
        set_cell_text(t.cell(r, 1), received, bold=True)
        set_cell_text(t.cell(r, 2), "")
        set_cell_text(t.cell(r, 3), "")
    add_callout(doc, "Decoder routine: 1 run checks -> 2 add failed labels -> 3 flip that position -> 4 extract 3,5,6,7", fill=PALE)

    add_title(doc, "Partner packet exchange", "Page 5 | Create, corrupt, swap, rescue", new_page=True)
    task_heading(doc, "5A", "Sender lab")
    doc.add_paragraph("Choose any four-bit message, encode it, have your partner audit it, then flip exactly one bit.")
    response_lines(doc, 1, "Message: __________________   Valid codeword: __________________")
    add_bit_grid(doc, bits=["", "", "", "", "", "", ""])
    response_lines(doc, 1, "Secret flipped position: ______   Corrupted transmission to send: __________________")

    task_heading(doc, "5B", "Receiver rescue")
    response_lines(doc, 1, "Received transmission: __________________")
    t = doc.add_table(rows=4, cols=4)
    set_table_geometry(t, [1800, 2600, 2280, 2680])
    for c, text in enumerate(["check", "count of 1s", "pass/fail", "failed label"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=9)
        shade_cell(t.cell(0, c), PALE2)
    for r, label in enumerate(["1", "2", "4"], start=1):
        set_cell_text(t.cell(r, 0), label, bold=True)
        for c in range(1, 4):
            set_cell_text(t.cell(r, c), "")
    response_lines(doc, 2, "Syndrome: ______   Corrected codeword: __________________\nRecovered message: __________________   Sender confirms? □")
    add_callout(doc, "Success means the corrected codeword, recovered message, and explanation all agree.", fill=GREEN_PALE, color=GREEN)

    add_title(doc, "Solo certification", "Page 6 | Notes allowed; no partner; show every check", new_page=True)
    task_heading(doc, "6A", "Part A - encode 1 0 0 1")
    add_bit_grid(doc, bits=["", "", "1", "", "0", "0", "1"])
    add_check_math(doc)
    response_lines(doc, 1, "Final codeword: __________________________________________")

    task_heading(doc, "6A", "Part B - repair 0 0 1 1 1 0 1", color=CORAL)
    add_bit_grid(doc, bits=["0", "0", "1", "1", "1", "0", "1"])
    response_lines(doc, 2, "Failed checks: __________________   Error position: ______\nCorrected codeword: __________________   Recovered message: __________________")

    task_heading(doc, "6B", "Bonus - extend to Hamming(8,4)")
    doc.add_paragraph("Add one overall even-parity bit as position 8. Complete the decision table.")
    t = doc.add_table(rows=5, cols=3)
    set_table_geometry(t, [3000, 3000, 3360])
    for c, text in enumerate(["Hamming syndrome", "overall parity", "what should the receiver do?"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=8)
        shade_cell(t.cell(0, c), PALE2)
    cases = [("0", "even"), ("nonzero", "odd"), ("0", "odd"), ("nonzero", "even")]
    for r, (syn, parity) in enumerate(cases, start=1):
        set_cell_text(t.cell(r, 0), syn)
        set_cell_text(t.cell(r, 1), parity)
        set_cell_text(t.cell(r, 2), "")
    response_lines(doc, 1, "Exit reflection: The step I am least likely to forget is ______________________________ because ______________________________")

    doc.core_properties.title = "Signal Rescue Lab - Hamming(7,4) Student Worksheet"
    doc.core_properties.subject = "Interactive grade 9-10 math circle worksheet"
    doc.core_properties.author = "Math Circle"
    doc.save(STUDENT)


def build_key():
    doc = Document()
    configure_document(doc, "SIGNAL RESCUE LAB  |  INSTRUCTOR ANSWER KEY")
    add_title(doc, "Signal Rescue Lab", "Instructor answer key | Hamming(7,4) | Even parity", key=True)
    add_callout(doc, "Use this key to check reasoning, not only final codewords. A correct syndrome without visible parity counts is not yet independent mastery.", fill=GREEN_PALE, color=GREEN)

    task_heading(doc, "1A", "The impossible repair")
    doc.add_paragraph("The missing bit cannot be determined from the damaged string alone. Both 0 and 1 are possible unless the sender adds an agreed rule or redundancy.")

    task_heading(doc, "1B", "Even-parity sprint")
    t = doc.add_table(rows=5, cols=4)
    set_table_geometry(t, [2600, 1900, 2260, 2600])
    for c, text in enumerate(["data", "check bit", "total 1s", "audit"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=9)
        shade_cell(t.cell(0, c), PALE2)
    rows = [("101", "0", "2", "even"), ("111", "1", "4", "even"), ("000", "0", "0", "even"), ("0101", "0", "2", "even")]
    for r, values in enumerate(rows, start=1):
        for c, value in enumerate(values):
            set_cell_text(t.cell(r, c), value, bold=c in (0, 1), color=GREEN if c == 3 else INK)

    task_heading(doc, "1C", "Address detective")
    add_coverage_grid(doc, answer=True)
    doc.add_paragraph("Addresses: position 3 = 1+2; position 5 = 1+4; position 6 = 2+4; position 7 = 1+2+4.")
    add_callout(doc, "Facilitator move: students should say which questions contain the position before adding the labels.", fill=PALE2)

    add_title(doc, "Seven-slot machine and warm-ups", "Answer key page 2", new_page=True, top_spacer=True)
    task_heading(doc, "2A", "Roles and check groups")
    add_bit_grid(doc, bits=["p1", "p2", "d1", "p4", "d2", "d3", "d4"])
    add_coverage_grid(doc, answer=True)

    task_heading(doc, "2B", "Training wheels: message 1 0 1")
    add_bit_grid(doc, bits=["1", "0", "1", "1", "0", "1", "0"], fixed7=True)
    add_check_math(doc, ["check 1: 1+1+0+0=2", "check 2: 0+1+1+0=2", "check 4: 1+0+1+0=2"])
    doc.add_paragraph("Completed word: 1011010.")

    task_heading(doc, "2C", "Training wheels: message 0 1 1")
    add_bit_grid(doc, bits=["1", "1", "0", "0", "1", "1", "0"], fixed7=True)
    doc.add_paragraph("Completed word: 1100110. Each parity group contains two 1s.")
    add_callout(doc, "Terminology: the warm-up fixes position 7 at zero. It is a restricted subcode of Hamming(7,4), not the standard Hamming code.", fill=AMBER_PALE, color=INK)

    add_title(doc, "Full Hamming(7,4) encoding", "Answer key page 3", new_page=True)
    task_heading(doc, "3A", "Four-move routine")
    doc.add_paragraph("1 label positions 1-7; 2 place data in 3,5,6,7; 3 choose p1,p2,p4 for even parity; 4 verify all three checks.")

    task_heading(doc, "3B", "Message 1 0 1 1")
    add_bit_grid(doc, bits=["0", "1", "1", "0", "0", "1", "1"])
    add_check_math(doc, ["check 1: 0+1+0+1=2", "check 2: 1+1+1+1=4", "check 4: 0+0+1+1=2"])
    doc.add_paragraph("Codeword: 0110011. Check bits: p1=0, p2=1, p4=0.")

    task_heading(doc, "3C", "Speed round")
    t = doc.add_table(rows=5, cols=3)
    set_table_geometry(t, [2400, 3480, 3480])
    for c, text in enumerate(["message", "codeword", "check bits p1,p2,p4"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=9)
        shade_cell(t.cell(0, c), PALE2)
    answers = [("0000", "0000000", "0,0,0"), ("1111", "1111111", "1,1,1"), ("0101", "0100101", "0,1,0"), ("1100", "0111100", "0,1,1")]
    for r, values in enumerate(answers, start=1):
        for c, value in enumerate(values):
            set_cell_text(t.cell(r, c), value, bold=c < 2)
    add_callout(doc, "Common error: placing message bits consecutively in positions 1-4. Require the role row before any parity arithmetic.", fill=CORAL_PALE, color=INK)

    add_title(doc, "Repair and decode", "Answer key page 4", new_page=True)
    task_heading(doc, "4A", "Repair 0 1 1 0 0 0 1")
    add_bit_grid(doc, bits=["0", "1", "1", "0", "0", "0", "1"], color_bad=6)
    t = doc.add_table(rows=4, cols=5)
    set_table_geometry(t, [1500, 3000, 1500, 1500, 1860])
    for c, text in enumerate(["check", "positions", "1s", "result", "label"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=8)
        shade_cell(t.cell(0, c), PALE2)
    rows = [("1", "1,3,5,7", "2", "pass", "0"), ("2", "2,3,6,7", "3", "fail", "2"), ("4", "4,5,6,7", "1", "fail", "4")]
    for r, values in enumerate(rows, start=1):
        for c, value in enumerate(values):
            set_cell_text(t.cell(r, c), value, bold=c in (0, 3), color=CORAL if value == "fail" else (GREEN if value == "pass" else INK))
    doc.add_paragraph("Syndrome = 2+4 = 6. Flip position 6. Corrected word = 0110011. Recovered message = 1011.")

    task_heading(doc, "4B", "Packet repair race")
    t = doc.add_table(rows=3, cols=5)
    set_table_geometry(t, [1000, 1900, 1760, 2350, 2350])
    for c, text in enumerate(["packet", "received", "failed", "corrected", "message"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=8)
        shade_cell(t.cell(0, c), PALE2)
    packets = [("A", "0101101", "4", "0100101", "0101"), ("B", "0111101", "1+2+4=7", "0111100", "1100")]
    for r, values in enumerate(packets, start=1):
        for c, value in enumerate(values):
            set_cell_text(t.cell(r, c), value, bold=c in (0, 2))
    add_callout(doc, "No-error case: if all checks pass, syndrome 0 means leave the codeword unchanged. Under the single-error promise, any syndrome 1-7 tells which bit to flip.", fill=GREEN_PALE, color=GREEN)

    add_title(doc, "Exchange, certification, and bonus", "Answer key page 5", new_page=True)
    task_heading(doc, "5A-5B", "Partner exchange")
    doc.add_paragraph("Answers vary. A valid sender word has even parity in all three groups. The receiver's syndrome must equal the sender's secretly flipped position, and extraction after repair must reproduce the sender's four-bit message.")
    add_callout(doc, "Extension: ask early finishers to flip a check bit (1,2,4) and compare with flipping a message bit (3,5,6,7). Both are corrected; only a data-bit error changes the extracted message before repair.", fill=PALE)

    task_heading(doc, "6A", "Solo certification")
    t = doc.add_table(rows=3, cols=4)
    set_table_geometry(t, [1300, 3000, 2500, 2560])
    headers = ["part", "result", "checks / syndrome", "final message"]
    for c, text in enumerate(headers):
        set_cell_text(t.cell(0, c), text, bold=True, size=8)
        shade_cell(t.cell(0, c), PALE2)
    answers = [("A", "0011001", "p1=0, p2=0, p4=1", "1001"), ("B", "error at 5; corrected 0011001", "failed 1+4", "1001")]
    for r, values in enumerate(answers, start=1):
        for c, value in enumerate(values):
            set_cell_text(t.cell(r, c), value, bold=c in (0, 1), color=GREEN if c == 3 else INK)

    task_heading(doc, "6B", "Extended Hamming(8,4) decision table")
    t = doc.add_table(rows=5, cols=3)
    set_table_geometry(t, [3000, 3000, 3360])
    for c, text in enumerate(["syndrome", "overall parity", "receiver action"]):
        set_cell_text(t.cell(0, c), text, bold=True, size=8)
        shade_cell(t.cell(0, c), PALE2)
    cases = [
        ("0", "even", "no error"),
        ("nonzero", "odd", "correct one error in positions 1-7"),
        ("0", "odd", "flip overall parity bit 8"),
        ("nonzero", "even", "detect two errors; do not correct"),
    ]
    for r, values in enumerate(cases, start=1):
        for c, value in enumerate(values):
            set_cell_text(t.cell(r, c), value, bold=c == 2)
    add_callout(doc, "Session mastery criterion: the student can encode and repair the solo certification without a partner, while showing the three check calculations.", fill=GREEN_PALE, color=GREEN)

    p = doc.add_paragraph("Source: R. W. Hamming, 'Error Detecting and Error Correcting Codes,' Bell System Technical Journal 29(2), 1950, pp. 147-160. https://doi.org/10.1002/j.1538-7305.1950.tb00463.x")
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    doc.core_properties.title = "Signal Rescue Lab - Hamming(7,4) Instructor Answer Key"
    doc.core_properties.subject = "Answer key and facilitation notes"
    doc.core_properties.author = "Math Circle"
    doc.save(KEY)


if __name__ == "__main__":
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_student()
    build_key()
    print(STUDENT)
    print(KEY)
